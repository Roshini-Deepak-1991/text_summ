# # # # import streamlit as st
# # # # import requests
# # # # import os
# # # # from dotenv import load_dotenv

# # # # load_dotenv()
# # # # HF_TOKEN = os.getenv("HF_TOKEN_TEXT")  # Changed variable name here

# # # # # load sentiment analysis pipeline
# # # # sentiment_pipeline = pipeline("sentiment-analysis")


# # # # def query(payload, api_url, headers):
# # # #     try:
# # # #         response = requests.post(api_url, headers=headers, json=payload)
# # # #         response.raise_for_status()
# # # #         return response.json()
# # # #     except requests.exceptions.RequestException as e:
# # # #         print(f"Raw API Response: {e.response.content if e.response else e}")
# # # #         return {"error": f"API request failed: {e}"}

# # # # def summarize_text(text, api_url, headers):
# # # #     output = query({"inputs": text}, api_url, headers)

# # # #     if "error" in output:
# # # #         return output["error"]

# # # #     try:
# # # #         return output[0]["summary_text"]
# # # #     except (KeyError, IndexError, TypeError) as e:
# # # #         return f"Error processing API response: {e}. Raw response: {output}"

# # # # def main():
# # # #     st.title("Feedback Summarizer (Hugging Face)")

# # # #     feedback_input = st.text_area("Enter feedback:", height=100)

# # # #     model = st.selectbox("Select Model", ["facebook/bart-large-cnn", "google/pegasus-xsum", "sshleifer/distilbart-cnn-12-6"])

# # # #     API_URL = f"https://api-inference.huggingface.co/models/{model}"

# # # #     headers = {"Authorization": f"Bearer {HF_TOKEN}"}  # Changed variable name here

# # # #     if st.button("Summarize"):
# # # #         if feedback_input:
# # # #             summary = summarize_text(feedback_input, API_URL, headers)
# # # #             sentiment, sentiment_score = analyze_sentiment(feedback_input)
# # # #             topics = identify_topics(feedback_input)

# # # #             st.subheader("Summary:")
# # # #             st.write(summary)
            
# # # #             st.subheader("Sentiment:")
# # # #             st.write(f"{sentiment} (Score: {sentiment_score:.2f})")

# # # #             st.subheader("Topics:")
# # # #             st.write(topics)

# # # #         else:
# # # #             st.warning("Please enter some feedback.")

# # # # if __name__ == "__main__":
# # # #     main()


# # # # # def query(payload, api_url, headers):
# # # # #     try:
# # # # #         response = requests.post(api_url, headers=headers, json=payload)
# # # # #         response.raise_for_status()
# # # # #         return response.json()
# # # # #     except requests.exceptions.RequestException as e:
# # # # #         print(f"Raw API Response: {e.response.content if e.response else e}")
# # # # #         return {"error": f"API request failed: {e}"}



# # # import streamlit as st
# # # import requests
# # # import os
# # # from dotenv import load_dotenv
# # # from transformers import pipeline

# # # load_dotenv()
# # # HF_TOKEN = os.getenv("HF_TOKEN_TEXT")

# # # sentiment_pipeline = pipeline("sentiment-analysis")
# # # keyword_extraction_pipeline = pipeline("text2text-generation", model="google/flan-t5-base") #Choose a suitable model

# # # def query(payload, api_url, headers):
# # #     try:
# # #         response = requests.post(api_url, headers=headers, json=payload)
# # #         response.raise_for_status()
# # #         return response.json()
# # #     except requests.exceptions.RequestException as e:
# # #         print(f"Raw API Response: {e.response.content if e.response else e}")
# # #         return {"error": f"API request failed: {e}"}

# # # def summarize_text(text, api_url, headers):
# # #     output = query({"inputs": text}, api_url, headers)
# # #     if "error" in output:
# # #         return output["error"]
# # #     try:
# # #         return output[0]["summary_text"]
# # #     except (KeyError, IndexError, TypeError) as e:
# # #         return f"Error processing API response: {e}. Raw response: {output}"

# # # def analyze_sentiment(text):
# # #     sentiment_result = sentiment_pipeline(text)[0]
# # #     return sentiment_result["label"], sentiment_result["score"]

# # # def extract_keywords(text):
# # #     prompt = f"extract keywords from: {text}"
# # #     keywords = keyword_extraction_pipeline(prompt, max_length=50)[0]['generated_text']
# # #     return keywords

# # # def main():
# # #     st.title("Feedback Analyzer (Hugging Face)")

# # #     feedback_input = st.text_area("Enter feedback:", height=100)

# # #     model = st.selectbox("Select Model", ["facebook/bart-large-cnn", "google/pegasus-xsum", "sshleifer/distilbart-cnn-12-6"])

# # #     API_URL = f"https://api-inference.huggingface.co/models/{model}"

# # #     headers = {"Authorization": f"Bearer {HF_TOKEN}"}

# # #     if st.button("Analyze"):
# # #         if feedback_input:
# # #             summary = summarize_text(feedback_input, API_URL, headers)
# # #             sentiment, sentiment_score = analyze_sentiment(feedback_input)
# # #             keywords = extract_keywords(feedback_input)

# # #             st.subheader("Summary:")
# # #             st.write(summary)

# # #             st.subheader("Sentiment:")
# # #             st.write(f"{sentiment} (Score: {sentiment_score:.2f})")

# # #             st.subheader("Keywords:")
# # #             st.write(keywords)
# # #         else:
# # #             st.warning("Please enter some feedback.")

# # # if __name__ == "__main__":
# # #     main()


# # import streamlit as st
# # import requests
# # import os
# # import time
# # from dotenv import load_dotenv
# # from transformers import pipeline

# # load_dotenv()
# # HF_TOKEN = os.getenv("HF_TOKEN_TEXT")

# # sentiment_pipeline = pipeline("sentiment-analysis")
# # keyword_extraction_pipeline = pipeline("text2text-generation", model="google/flan-t5-base")

# # def query(payload, api_url, headers):
# #     max_retries = 3
# #     retry_delay = 2

# #     for retry in range(max_retries):
# #         try:
# #             response = requests.post(api_url, headers=headers, json=payload)
# #             response.raise_for_status()
# #             return response.json()
# #         except requests.exceptions.HTTPError as e:
# #             if e.response.status_code == 503 and retry < max_retries - 1:
# #                 st.warning(f"Service unavailable. Retrying in {retry_delay} seconds...")
# #                 time.sleep(retry_delay)
# #                 retry_delay *= 2
# #             else:
# #                 print(f"Raw API Response: {e.response.content if e.response else e}")
# #                 return {"error": f"API request failed: {e}"}
# #         except requests.exceptions.RequestException as e:
# #             print(f"Raw API Response: {e.response.content if e.response else e}")
# #             return {"error": f"API request failed: {e}"}

# # def summarize_text(text, api_url, headers):
# #     output = query({"inputs": text}, api_url, headers)
# #     if "error" in output:
# #         return output["error"]
# #     try:
# #         return output[0]["summary_text"]
# #     except (KeyError, IndexError, TypeError) as e:
# #         return f"Error processing API response: {e}. Raw response: {output}"

# # def analyze_sentiment(text):
# #     sentiment_result = sentiment_pipeline(text)[0]
# #     return sentiment_result["label"], sentiment_result["score"]

# # def extract_keywords(text):
# #     prompt = f"extract keywords from: {text}"
# #     keywords = keyword_extraction_pipeline(prompt, max_length=50)[0]['generated_text']
# #     return keywords

# # def main():
# #     st.title("Feedback Analyzer (Hugging Face)")

# #     feedback_input = st.text_area("Enter feedback:", height=100)

# #     models = ["facebook/bart-large-cnn", "google/pegasus-xsum", ]
# #     selected_model = st.selectbox("Select Model", models)
# #     API_URL = f"https://api-inference.huggingface.co/models/{selected_model}"
# #     headers = {"Authorization": f"Bearer {HF_TOKEN}"}

# #     if st.button("Analyze"):
# #         if feedback_input:
# #             summary = summarize_text(feedback_input, API_URL, headers)
# #             if "error" in summary: #If the first model fails.
# #                 st.warning(f"Model {selected_model} failed. Trying fallback.")
# #                 for fallback_model in models: #Iterate through models.
# #                     if fallback_model != selected_model: #Ensure not the same model.
# #                         fallback_url = f"https://api-inference.huggingface.co/models/{fallback_model}"
# #                         summary = summarize_text(feedback_input, fallback_url, headers)
# #                         if "error" not in summary: #If fallback works, break.
# #                             st.info(f"Fallback model {fallback_model} used.")
# #                             break
# #             sentiment, sentiment_score = analyze_sentiment(feedback_input)
# #             keywords = extract_keywords(feedback_input)

# #             st.subheader("Summary:")
# #             st.write(summary)

# #             st.subheader("Sentiment:")
# #             st.write(f"{sentiment} (Score: {sentiment_score:.2f})")

# #             st.subheader("Keywords:")
# #             st.write(keywords)
# #         else:
# #             st.warning("Please enter some feedback.")

# # if __name__ == "__main__":
# #     main()


# import streamlit as st
# import requests
# import os
# import time
# from dotenv import load_dotenv
# from transformers import pipeline
# import PyPDF2
# from docx import Document

# load_dotenv()
# HF_TOKEN = os.getenv("HF_TOKEN_TEXT")

# sentiment_pipeline = pipeline("sentiment-analysis")
# keyword_extraction_pipeline = pipeline("text2text-generation", model="google/flan-t5-base")

# def query(payload, api_url, headers):
#     max_retries = 3
#     retry_delay = 2

#     for retry in range(max_retries):
#         try:
#             response = requests.post(api_url, headers=headers, json=payload)
#             response.raise_for_status()
#             return response.json()
#         except requests.exceptions.HTTPError as e:
#             if e.response.status_code == 503 and retry < max_retries - 1:
#                 st.warning(f"Service unavailable. Retrying in {retry_delay} seconds...")
#                 time.sleep(retry_delay)
#                 retry_delay *= 2
#             else:
#                 print(f"Raw API Response: {e.response.content if e.response else e}")
#                 return {"error": f"API request failed: {e}"}
#         except requests.exceptions.RequestException as e:
#             print(f"Raw API Response: {e.response.content if e.response else e}")
#             return {"error": f"API request failed: {e}"}

# def summarize_text(text, api_url, headers):
#     output = query({"inputs": text}, api_url, headers)
#     if "error" in output:
#         return output["error"]
#     try:
#         return output[0]["summary_text"]
#     except (KeyError, IndexError, TypeError) as e:
#         return f"Error processing API response: {e}. Raw response: {output}"

# def analyze_sentiment(text):
#     sentiment_result = sentiment_pipeline(text)[0]
#     return sentiment_result["label"], sentiment_result["score"]

# def extract_keywords(text):
#     prompt = f"extract keywords from: {text}"
#     keywords = keyword_extraction_pipeline(prompt, max_length=50)[0]['generated_text']
#     return keywords

# def read_pdf(file):
#     pdf_reader = PyPDF2.PdfReader(file)
#     text = ""
#     for page in pdf_reader.pages:
#         text += page.extract_text()
#     return text

# def read_docx(file):
#     doc = Document(file)
#     text = ""
#     for paragraph in doc.paragraphs:
#         text += paragraph.text + "\n"
#     return text

# def main():
#     st.title("Feedback Analyzer (Hugging Face)")

#     uploaded_file = st.file_uploader("Upload PDF, DOCX, or TXT file (Max 10MB)", type=["pdf", "docx", "txt"])

#     if uploaded_file is not None:
#         file_size_mb = uploaded_file.size / (1024 * 1024)
#         if file_size_mb > 10:
#             st.error("File size exceeds 10MB limit.")
#             return

#         file_extension = uploaded_file.name.split(".")[-1].lower()
#         if file_extension == "pdf":
#             text = read_pdf(uploaded_file)
#         elif file_extension == "docx":
#             text = read_docx(uploaded_file)
#         elif file_extension == "txt":
#             text = uploaded_file.getvalue().decode("utf-8")
#         else:
#             st.error("Unsupported file type.")
#             return

#         feedback_input = st.text_area("Extracted Text:", value=text, height=200)
#     else:
#         feedback_input = st.text_area("Enter feedback:", height=100)

#     models = ["facebook/bart-large-cnn", "google/pegasus-xsum", "sshleifer/distilbart-cnn-12-6"]
#     selected_model = st.selectbox("Select Model", models)
#     API_URL = f"https://api-inference.huggingface.co/models/{selected_model}"
#     headers = {"Authorization": f"Bearer {HF_TOKEN}"}

#     if st.button("Analyze"):
#         if feedback_input:
#             summary = summarize_text(feedback_input, API_URL, headers)
#             if "error" in summary:
#                 st.warning(f"Model {selected_model} failed. Trying fallback.")
#                 for fallback_model in models:
#                     if fallback_model != selected_model:
#                         fallback_url = f"https://api-inference.huggingface.co/models/{fallback_model}"
#                         summary = summarize_text(feedback_input, fallback_url, headers)
#                         if "error" not in summary:
#                             st.info(f"Fallback model {fallback_model} used.")
#                             break
#             sentiment, sentiment_score = analyze_sentiment(feedback_input)
#             keywords = extract_keywords(feedback_input)

#             st.subheader("Summary:")
#             st.write(summary)

#             st.subheader("Sentiment:")
#             st.write(f"{sentiment} (Score: {sentiment_score:.2f})")

#             st.subheader("Keywords:")
#             st.write(keywords)
#         else:
#             st.warning("Please enter some feedback or upload a file.")

# if __name__ == "__main__":
#     main()

import streamlit as st
import requests
import os
import time
import asyncio
import httpx
from dotenv import load_dotenv
from transformers import pipeline
import PyPDF2
from docx import Document

load_dotenv()
HF_TOKEN = os.getenv("HF_TOKEN_TEXT")

sentiment_pipeline = pipeline("sentiment-analysis")
keyword_extraction_pipeline = pipeline("text2text-generation", model="google/flan-t5-base")

async def async_query(payload, api_url, headers):
    async with httpx.AsyncClient() as client:
        try:
            response = await client.post(api_url, headers=headers, json=payload)
            response.raise_for_status()
            return response.json()
        except httpx.HTTPError as e:
            if e.response is not None and e.response.status_code == 503:
                return {"error": f"Service unavailable: {e}"}
            else:
                return {"error": f"API request failed: {e}"}
        except httpx.RequestError as e:
            return {"error": f"API request failed: {e}"}

async def summarize_text(text, api_url, headers):
    output = await async_query({"inputs": text}, api_url, headers)
    if "error" in output:
        return output["error"]
    try:
        return output[0]["summary_text"]
    except (KeyError, IndexError, TypeError) as e:
        return f"Error processing API response: {e}. Raw response: {output}"

def analyze_sentiment(text):
    sentiment_result = sentiment_pipeline(text)[0]
    return sentiment_result["label"], sentiment_result["score"]

def extract_keywords(text):
    prompt = f"extract keywords from: {text}"
    keywords = keyword_extraction_pipeline(prompt, max_length=50)[0]['generated_text']
    return keywords

def read_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def read_docx(file):
    doc = Document(file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

async def main():
    st.title("Feedback Analyzer (Hugging Face)")

    uploaded_file = st.file_uploader("Upload PDF, DOCX, or TXT file (Max 10MB)", type=["pdf", "docx", "txt"])

    if uploaded_file is not None:
        file_size_mb = uploaded_file.size / (1024 * 1024)
        if file_size_mb > 10:
            st.error("File size exceeds 10MB limit.")
            return

        file_extension = uploaded_file.name.split(".")[-1].lower()
        if file_extension == "pdf":
            text = read_pdf(uploaded_file)
        elif file_extension == "docx":
            text = read_docx(uploaded_file)
        elif file_extension == "txt":
            text = uploaded_file.getvalue().decode("utf-8")
        else:
            st.error("Unsupported file type.")
            return

        feedback_input = st.text_area("Extracted Text:", value=text, height=200)
    else:
        feedback_input = st.text_area("Enter feedback:", height=100)

    models = ["facebook/bart-large-cnn", "google/pegasus-xsum", "sshleifer/distilbart-cnn-12-6"]
    selected_model = st.selectbox("Select Model", models)
    API_URL = f"https://api-inference.huggingface.co/models/{selected_model}"
    headers = {"Authorization": f"Bearer {HF_TOKEN}"}

    if st.button("Analyze"):
        if feedback_input:
            summary = await summarize_text(feedback_input, API_URL, headers)
            if "error" in summary:
                st.warning(f"Model {selected_model} failed. Trying fallback.")
                for fallback_model in models:
                    if fallback_model != selected_model:
                        fallback_url = f"https://api-inference.huggingface.co/models/{fallback_model}"
                        summary = await summarize_text(feedback_input, fallback_url, headers)
                        if "error" not in summary:
                            st.info(f"Fallback model {fallback_model} used.")
                            break

            sentiment, sentiment_score = analyze_sentiment(feedback_input)
            keywords = extract_keywords(feedback_input)

            st.subheader("Summary:")
            st.write(summary)

            st.subheader("Sentiment:")
            st.write(f"{sentiment} (Score: {sentiment_score:.2f})")

            st.subheader("Keywords:")
            st.write(keywords)
        else:
            st.warning("Please enter some feedback or upload a file.")

if __name__ == "__main__":
    asyncio.run(main())